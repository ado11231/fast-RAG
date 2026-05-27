"""Tests for DocxLoader — DOCX text extraction via python-docx."""
from pathlib import Path

import pytest

from fastrag.loaders.docx import DocxLoader


def test_docx_loader_returns_string(docx_sample: Path) -> None:
    """Verify the loader extracts text from a basic .docx file."""
    loader = DocxLoader()
    result = loader.load(docx_sample)
    assert isinstance(result, str)
    assert len(result) > 0


def test_docx_loader_contains_expected_text(docx_sample: Path) -> None:
    """The extracted text should contain the words we wrote."""
    loader = DocxLoader()
    result = loader.load(docx_sample)
    assert "Hello" in result


@pytest.fixture
def docx_sample(tmp_path: Path) -> Path:
    """Generate a minimal .docx with two paragraphs for testing."""
    path = tmp_path / "sample.docx"
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello from DOCX")
        doc.add_paragraph("Second paragraph")
        doc.save(str(path))
    except ImportError:
        pytest.skip("python-docx not installed — cannot generate DOCX fixture")
    return path
